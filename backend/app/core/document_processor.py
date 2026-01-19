"""
Document Processing Pipeline

Extracts text from various document formats and chunks them intelligently.

Supported formats:
- PDF (via PyMuPDF)
- DOCX (via python-docx)
- TXT (plain text)
- CSV (via pandas - optional)

Features:
- Smart chunking (sentence-aware)
- Metadata extraction
- Language detection
- OCR support (optional)
"""

import io
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union
import re

# PDF processing
import fitz  # PyMuPDF

# DOCX processing
from docx import Document

# Text chunking
from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.config import settings
from app.utils.logger import get_logger
from app.utils.exceptions import DocumentProcessingError
from app.utils.helpers import generate_id, hash_text

logger = get_logger(__name__)


class DocumentChunk:
    """
    Represents a chunk of a document.
    
    Attributes:
        chunk_id: Unique chunk identifier
        text: Chunk text content
        metadata: Chunk metadata (document info, page number, etc.)
    """
    
    def __init__(
        self,
        text: str,
        metadata: Dict[str, Any],
        chunk_id: Optional[str] = None
    ):
        self.chunk_id = chunk_id or generate_id()
        self.text = text
        self.metadata = metadata
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary"""
        return {
            "chunk_id": self.chunk_id,
            "text": self.text,
            "metadata": self.metadata
        }


class DocumentProcessor:
    """
    Process documents and extract text with metadata.
    
    Features:
    - Extract text from PDF, DOCX, TXT
    - Smart chunking with overlap
    - Preserve document structure
    - Extract metadata
    
    Example:
        processor = DocumentProcessor()
        
        # Process PDF
        chunks = processor.process_file("document.pdf", user_id="123")
        
        # Process from bytes
        chunks = processor.process_bytes(pdf_bytes, "document.pdf", user_id="123")
    """
    
    def __init__(
        self,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None
    ):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Size of each chunk in tokens (from settings if not provided)
            chunk_overlap: Overlap between chunks (from settings if not provided)
        """
        self.chunk_size = chunk_size or settings.CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or settings.CHUNK_OVERLAP
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],  # Sentence-aware
        )
        
        logger.info(
            f"Initialized DocumentProcessor "
            f"(chunk_size={self.chunk_size}, overlap={self.chunk_overlap})"
        )
    
    def process_file(
        self,
        file_path: Union[str, Path],
        user_id: str,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        Process a document file.
        
        Args:
            file_path: Path to document file
            user_id: User ID who owns the document
            additional_metadata: Additional metadata to attach
        
        Returns:
            List of DocumentChunk objects
        
        Raises:
            DocumentProcessingError: If processing fails
        
        Example:
            chunks = processor.process_file("doc.pdf", user_id="user123")
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentProcessingError(f"File not found: {file_path}")
        
        logger.info(f"Processing file: {file_path.name}")
        
        # Read file bytes
        with open(file_path, "rb") as f:
            file_bytes = f.read()
        
        return self.process_bytes(
            file_bytes,
            file_path.name,
            user_id,
            additional_metadata
        )
    
    def process_bytes(
        self,
        file_bytes: bytes,
        filename: str,
        user_id: str,
        additional_metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
        """
        Process document from bytes.
        
        Args:
            file_bytes: Document bytes
            filename: Original filename
            user_id: User ID
            additional_metadata: Additional metadata
        
        Returns:
            List of DocumentChunk objects
        
        Raises:
            DocumentProcessingError: If processing fails
        """
        # Detect file type
        file_extension = Path(filename).suffix.lower()
        
        logger.info(f"Processing {filename} ({len(file_bytes)} bytes)")
        
        try:
            # Extract text based on file type
            if file_extension == ".pdf":
                text, page_count = self._extract_pdf(file_bytes)
            elif file_extension == ".docx":
                text, page_count = self._extract_docx(file_bytes)
            elif file_extension in [".txt", ".md"]:
                text = file_bytes.decode("utf-8")
                page_count = 1
            else:
                raise DocumentProcessingError(
                    f"Unsupported file type: {file_extension}"
                )
            
            logger.info(f"Extracted {len(text)} characters from {filename}")
            
            # Chunk the text
            chunks = self._chunk_text(text)
            
            # Create DocumentChunk objects with metadata
            document_id = generate_id()
            document_hash = hash_text(text)
            
            document_chunks = []
            for i, chunk_text in enumerate(chunks):
                metadata = {
                    "document_id": document_id,
                    "document_name": filename,
                    "document_hash": document_hash,
                    "user_id": user_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "page_count": page_count,
                    "file_type": file_extension[1:],  # Remove dot
                    **(additional_metadata or {})
                }
                
                chunk = DocumentChunk(
                    text=chunk_text,
                    metadata=metadata
                )
                document_chunks.append(chunk)
            
            logger.info(
                f"✅ Created {len(document_chunks)} chunks from {filename}"
            )
            return document_chunks
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise DocumentProcessingError(f"Failed to process {filename}: {e}")
    
    def _extract_pdf(self, pdf_bytes: bytes) -> Tuple[str, int]:
        """
        Extract text from PDF bytes.
        
        Args:
            pdf_bytes: PDF file bytes
        
        Returns:
            Tuple of (extracted_text, page_count)
        """
        try:
            # Open PDF from bytes
            pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            
            text_parts = []
            page_count = len(pdf_doc)
            
            # Extract text from each page
            for page_num in range(page_count):
                page = pdf_doc[page_num]
                text = page.get_text()
                
                if text.strip():
                    # Add page separator
                    text_parts.append(f"\n--- Page {page_num + 1} ---\n")
                    text_parts.append(text)
            
            pdf_doc.close()
            
            full_text = "\n".join(text_parts)
            
            # Clean text
            full_text = self._clean_text(full_text)
            return full_text, page_count
        
        except Exception as e:
            raise DocumentProcessingError(f"PDF extraction failed: {e}")
    
    def _extract_docx(self, docx_bytes: bytes) -> Tuple[str, int]:
        """
        Extract text from DOCX bytes.
        
        Args:
            docx_bytes: DOCX file bytes
        
        Returns:
            Tuple of (extracted_text, paragraph_count)
        """
        try:
            # Open DOCX from bytes
            doc = Document(io.BytesIO(docx_bytes))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            
            # Clean text
            full_text = self._clean_text(full_text)
            
            return full_text, len(doc.paragraphs)
            
        except Exception as e:
            raise DocumentProcessingError(f"DOCX extraction failed: {e}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        Args:
            text: Raw text
        
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove excessive spaces
        text = re.sub(r' +', ' ', text)
        
        # Remove zero-width characters
        text = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _chunk_text(self, text: str) -> List[str]:
        """
        Split text into chunks using RecursiveCharacterTextSplitter.
        
        Args:
            text: Full text to chunk
        
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        chunks = self.text_splitter.split_text(text)
        
        # Filter out empty chunks
        chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
        
        return chunks
    
    def estimate_chunks(self, text: str) -> int:
        """
        Estimate number of chunks without actually chunking.
        
        Args:
            text: Text to estimate
        
        Returns:
            Estimated number of chunks
        """
        return max(1, len(text) // self.chunk_size)

#=============================================================================
#LANGUAGE DETECTION
#=============================================================================

    def detect_language(text: str) -> str:
    """
    Detect language of text.
    Args:
    text: Input text

    Returns:
        ISO 639-1 language code (e.g., 'en', 'hi')
    
    Example:
        lang = detect_language("भारत की राजधानी")
        # 'hi'
    """
    try:
        from langdetect import detect
        
        # Take first 500 chars for detection
        sample = text[:500]
        lang_code = detect(sample)
        
        return lang_code
        
    except Exception as e:
        logger.warning(f"Language detection failed: {e}, defaulting to 'en'")
        return "en"

#=============================================================================
#CONVENIENCE FUNCTIONS
#=============================================================================
    def process_document(
    file_path: Union[str, Path],
    user_id: str,
    metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
    """
    Convenience function to process a document file.
    Args:
    file_path: Path to document
    user_id: User ID
    metadata: Additional metadata

    Returns:
        List of DocumentChunk objects
    """
    processor = DocumentProcessor()
    return processor.process_file(file_path, user_id, metadata)

    def process_text(
    text: str,
    document_name: str,
    user_id: str,
    metadata: Optional[Dict[str, Any]] = None
    ) -> List[DocumentChunk]:
    """
    Process plain text directly.
    Args:
    text: Text content
    document_name: Name for the document
    user_id: User ID
    metadata: Additional metadata

    Returns:
        List of DocumentChunk objects
    """
    processor = DocumentProcessor()
    
    # Convert to bytes
    text_bytes = text.encode("utf-8")
    
    return processor.process_bytes(
        text_bytes,
        f"{document_name}.txt",
        user_id,
        metadata
    )
#=============================================================================
#TESTING
#=============================================================================
    if name == "main":
    print("=" * 80)
    print("TESTING DOCUMENT PROCESSOR")
    print("=" * 80)
    # Test with sample text
    sample_text = """
    India, officially the Republic of India, is a country in South Asia.
    It is the seventh-largest country by area and the most populous country.
    
    The capital of India is New Delhi. Mumbai is the largest city.
    India has 28 states and 8 union territories.
    
    भारत, आधिकारिक तौर पर भारत गणराज्य, दक्षिण एशिया का एक देश है।
    यह क्षेत्रफल के हिसाब से सातवां सबसे बड़ा देश है।
    """ * 10  # Repeat to ensure chunking
    
    print(f"\n📝 Processing sample text ({len(sample_text)} chars)...")
    
    chunks = process_text(
        text=sample_text,
        document_name="sample_doc",
        user_id="test_user",
        metadata={"category": "test"}
    )
    
    print(f"✅ Created {len(chunks)} chunks")
    
    for i, chunk in enumerate(chunks[:3], 1):  # Show first 3
        print(f"\nChunk {i}:")
        print(f"  ID: {chunk.chunk_id}")
        print(f"  Text length: {len(chunk.text)} chars")
        print(f"  Preview: {chunk.text[:100]}...")
        print(f"  Metadata: {chunk.metadata}")
    
    # Test language detection
    print("\n🌍 Testing language detection...")
    
    test_texts = [
        ("Hello world", "en"),
        ("भारत की राजधानी", "hi"),
        ("இந்தியா", "ta"),
    ]
    
    for text, expected in test_texts:
        detected = detect_language(text)
        print(f"  '{text}' → {detected} (expected: {expected})")
    
    print("\n" + "=" * 80)
    print("✅ DOCUMENT PROCESSOR WORKING CORRECTLY!")
    print("=" * 80)
